# -*- coding: utf-8 -*-
"""
report_generator.py · 报告导出系统

提供三种格式导出（均带中文支持）：
    Excel —— 4 个 Sheet：原始数据 / 统计结果 / 交叉分析 / AI洞察
    PDF   —— 封面 / 目录 / 统计图 / 分析结果 / AI报告 / 建议措施
    Word  —— 同 PDF 结构

并提供：
    generate_sample_data() —— 一键生成 200 份模拟答卷用于展示；
    ReportPanel            —— 报告与示例数据的图形操作面板。

图表用 matplotlib 渲染为 PNG 后嵌入；PDF 使用 reportlab 内置中文字体 STSong-Light。
"""
from __future__ import annotations

import logging
import tkinter as tk
from datetime import datetime
from pathlib import Path
from tkinter import messagebox, scrolledtext
from typing import Any, Dict, List, Optional

from modules import data_manager as dm
from modules.analytics import Analyzer
from modules.ai_engine import AIEngine
from modules.questionnaire_engine import Questionnaire

log = logging.getLogger("report_generator")

BASE_DIR = Path(__file__).resolve().parent.parent
OPTION_TYPES = {"single", "multiple", "yesno"}
NUMERIC_TYPES = {"number", "rating"}


# ======================================================================
# 通用工具
# ======================================================================
def _exports_dir(config: Dict[str, Any], base_dir: Path) -> Path:
    d = base_dir / config.get("paths", {}).get("exports", "exports")
    d.mkdir(parents=True, exist_ok=True)
    return d


def _ts() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def build_ai_report(config: Dict[str, Any], q: Questionnaire, analyzer: Analyzer) -> Dict[str, Any]:
    """统一生成 AI 分析报告（离线/联网由 ai_engine 决定）。"""
    return AIEngine(config, q, analyzer).generate_report()


# ----------------------------------------------------------------------
# 图表渲染（matplotlib → PNG，浅色报告风格）
# ----------------------------------------------------------------------
def _setup_mpl():
    import matplotlib
    import matplotlib.pyplot as plt
    plt.rcParams["font.sans-serif"] = [
        "Arial Unicode MS", "PingFang SC", "Heiti TC", "STHeiti",
        "Microsoft YaHei", "SimHei", "WenQuanYi Zen Hei", "Noto Sans CJK SC",
    ]
    plt.rcParams["axes.unicode_minus"] = False
    return plt


def _pick(questions: List[Dict[str, Any]], types, n=1):
    return [q for q in questions if q.get("type") in types][:n]


def render_charts(q: Questionnaire, analyzer: Analyzer, outdir: Path) -> Dict[str, Path]:
    """渲染报告所需图表，返回 {名称: png路径}。matplotlib 缺失则返回空。"""
    try:
        plt = _setup_mpl()
    except Exception as exc:  # noqa: BLE001
        log.warning("matplotlib 不可用，报告将不含图表：%s", exc)
        return {}
    outdir.mkdir(parents=True, exist_ok=True)
    charts: Dict[str, Path] = {}
    palette = ["#2563eb", "#0ea5a4", "#f59e0b", "#ef4444", "#8b5cf6", "#ec4899", "#22c55e"]

    def save(fig, name) -> None:
        p = outdir / f"{name}.png"
        fig.savefig(p, dpi=130, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        charts[name] = p

    # 前两个选项题饼图
    for i, qd in enumerate(_pick(q.questions, {"single", "yesno"}, n=2)):
        st = analyzer.question_stats(qd["id"])
        dist = [d for d in st.get("distribution", []) if d["count"] > 0]
        if not dist:
            continue
        fig, ax = plt.subplots(figsize=(4.2, 3.4))
        ax.pie([d["count"] for d in dist], labels=[d["label"] for d in dist],
               autopct="%1.0f%%", colors=palette, textprops={"fontsize": 9})
        ax.set_title(qd.get("title", "")[:18], fontsize=12)
        save(fig, f"pie_{i}")

    # 首个评分题柱状
    for qd in _pick(q.questions, {"rating"}, n=1):
        st = analyzer.question_stats(qd["id"])
        dist = st.get("distribution", [])
        if dist:
            fig, ax = plt.subplots(figsize=(5.2, 3.2))
            ax.bar([d["label"] for d in dist], [d["count"] for d in dist], color="#2563eb")
            ax.set_title(qd.get("title", "")[:18] + f"（均值{st.get('mean','-')}）", fontsize=12)
            save(fig, "rating_bar")

    # 首个多选题排行
    for qd in _pick(q.questions, {"multiple"}, n=1):
        st = analyzer.question_stats(qd["id"])
        dist = sorted([d for d in st.get("distribution", []) if d["count"] > 0], key=lambda x: x["count"])
        if dist:
            fig, ax = plt.subplots(figsize=(5.2, 3.2))
            ax.barh([d["label"] for d in dist], [d["count"] for d in dist], color="#0ea5a4")
            ax.set_title(qd.get("title", "")[:18] + "（排行）", fontsize=12)
            save(fig, "multi_rank")

    # 趋势折线
    by_date = analyzer.count_by_date()
    if by_date:
        fig, ax = plt.subplots(figsize=(7.2, 3.0))
        ax.plot(list(by_date.keys()), list(by_date.values()), marker="o", color="#2563eb")
        ax.set_title("每日答卷数量趋势", fontsize=12)
        fig.autofmt_xdate(rotation=30)
        save(fig, "trend")

    return charts


# ======================================================================
# Excel 导出（4 Sheet）
# ======================================================================
def export_excel(config: Dict[str, Any], q: Questionnaire, analyzer: Analyzer,
                 base_dir: Path = BASE_DIR, ai_report: Optional[Dict[str, Any]] = None) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    analyzer.refresh()
    ai_report = ai_report or build_ai_report(config, q, analyzer)
    header_fill = PatternFill("solid", fgColor="2563EB")
    sub_fill = PatternFill("solid", fgColor="DBE4F5")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    bold = Font(bold=True)

    def style_header(ws, row=1):
        for cell in ws[row]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")

    wb = Workbook()

    # ---- Sheet1 原始数据 ----
    ws1 = wb.active
    ws1.title = "原始数据"
    headers = ["编号", "开始时间", "结束时间", "用时(秒)", "完成"] + \
              [f"{qq['id']} {qq.get('title','')[:14]}" for qq in q.questions]
    ws1.append(headers)
    style_header(ws1)
    for r in analyzer.records:
        row = [r.get("response_id"), r.get("started_at"), r.get("finished_at"),
               r.get("duration_seconds"), "是" if r.get("completed") else "否"]
        for qq in q.questions:
            v = r.get("answers", {}).get(qq["id"], "")
            row.append(q.answer_display(qq["id"], v) if v != "" else "")
        ws1.append(row)
    ws1.freeze_panes = "A2"
    for i, _ in enumerate(headers, 1):
        ws1.column_dimensions[get_column_letter(i)].width = 16

    # ---- Sheet2 统计结果 ----
    ws2 = wb.create_sheet("统计结果")
    ws2.append(["题目", "选项 / 指标", "计数 / 数值", "占比"])
    style_header(ws2)
    for st in analyzer.all_question_stats():
        title = f"{st['id']} {st['title'][:20]}"
        if st.get("type") in OPTION_TYPES:
            for d in st.get("distribution", []):
                ws2.append([title, d["label"], d["count"], f"{d['ratio']*100:.1f}%"])
        elif st.get("type") in NUMERIC_TYPES and st.get("n"):
            for k, label in [("mean", "均值"), ("median", "中位数"), ("std", "标准差"),
                             ("min", "最小值"), ("max", "最大值"), ("n", "样本数")]:
                ws2.append([title, label, st.get(k), ""])
        elif st.get("type") == "text":
            ws2.append([title, "填写条数", st.get("filled", 0), ""])
        ws2[f"A{ws2.max_row}"].font = bold
    for col, w in zip("ABCD", (34, 22, 14, 10)):
        ws2.column_dimensions[col].width = w

    # ---- Sheet3 交叉分析 ----
    ws3 = wb.create_sheet("交叉分析")
    rownum = 1
    for a, b in analyzer.suggested_crosstabs():
        ct = analyzer.crosstab(a, b)
        if not ct or ct["total"] == 0:
            continue
        ws3.cell(rownum, 1, f"{ct['row_title'][:18]}  ×  {ct['col_title'][:18]}（n={ct['total']}）").font = bold
        rownum += 1
        # 表头
        ws3.cell(rownum, 1, "")
        for j, cl in enumerate(ct["col_labels"], 2):
            c = ws3.cell(rownum, j, cl); c.font = bold; c.fill = sub_fill
        ws3.cell(rownum, len(ct["col_labels"]) + 2, "合计").font = bold
        rownum += 1
        for i, rl in enumerate(ct["row_labels"]):
            c = ws3.cell(rownum, 1, rl); c.font = bold; c.fill = sub_fill
            for j, _ in enumerate(ct["col_labels"]):
                ws3.cell(rownum, j + 2, ct["matrix"][i][j])
            ws3.cell(rownum, len(ct["col_labels"]) + 2, ct["row_totals"][i])
            rownum += 1
        rownum += 1  # 空行分隔
    if rownum == 1:
        ws3.cell(1, 1, "（暂无足够数据进行交叉分析）")
    for col in "ABCDEFGH":
        ws3.column_dimensions[col].width = 14

    # ---- Sheet4 AI洞察 ----
    ws4 = wb.create_sheet("AI洞察")
    ws4.append(["类别", "内容"])
    style_header(ws4)
    sections = [("调查结论", "conclusions"), ("用户画像", "persona"), ("趋势预测", "trends"),
                ("改进建议", "suggestions"), ("风险提示", "risks"), ("智能洞察", "insights")]
    for label, key in sections:
        for item in ai_report.get(key, []):
            ws4.append([label, item])
            ws4[f"A{ws4.max_row}"].font = bold
    ws4.column_dimensions["A"].width = 12
    ws4.column_dimensions["B"].width = 90
    for row in ws4.iter_rows(min_row=2):
        row[1].alignment = Alignment(wrap_text=True, vertical="top")

    out = _exports_dir(config, base_dir) / f"survey_report_{_ts()}.xlsx"
    wb.save(out)
    log.info("已导出 Excel：%s", out)
    return out


# ======================================================================
# PDF 导出
# ======================================================================
def export_pdf(config: Dict[str, Any], q: Questionnaire, analyzer: Analyzer,
               base_dir: Path = BASE_DIR, ai_report: Optional[Dict[str, Any]] = None) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import (Image, PageBreak, Paragraph, SimpleDocTemplate,
                                    Spacer, Table, TableStyle)

    analyzer.refresh()
    ai_report = ai_report or build_ai_report(config, q, analyzer)
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))  # reportlab 内置中文字体
    FONT = "STSong-Light"

    base = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=base["Heading1"], fontName=FONT, fontSize=20, leading=28, spaceAfter=12)
    h2 = ParagraphStyle("h2", parent=base["Heading2"], fontName=FONT, fontSize=15, leading=22,
                        textColor=colors.HexColor("#1d4ed8"), spaceBefore=14, spaceAfter=8)
    body = ParagraphStyle("body", parent=base["Normal"], fontName=FONT, fontSize=11, leading=18)
    cover_title = ParagraphStyle("ct", parent=base["Title"], fontName=FONT, fontSize=26, leading=36)
    small = ParagraphStyle("small", parent=base["Normal"], fontName=FONT, fontSize=10,
                          textColor=colors.grey, alignment=1)

    charts = render_charts(q, analyzer, _exports_dir(config, base_dir) / f"_charts_{_ts()}")
    s = ai_report["summary"]
    story: List[Any] = []

    # ---- 封面 ----
    story += [Spacer(1, 5 * cm),
              Paragraph(config.get("export", {}).get("report_title", "适老化智能问卷调查分析报告"), cover_title),
              Spacer(1, 1 * cm),
              Paragraph(q.title(), h2),
              Spacer(1, 3 * cm),
              Paragraph(f"样本规模：{s['total']} 份　|　完成率：{s['completion_rate']*100:.1f}%", body),
              Paragraph(f"生成时间：{ai_report['generated_at']}　|　分析模式：{'联网AI' if ai_report['mode']=='online' else '离线规则'}", body),
              Spacer(1, 4 * cm),
              Paragraph(config.get("export", {}).get("company_name", "适老化调研课题组"), small),
              PageBreak()]

    # ---- 目录 ----
    story += [Paragraph("目  录", h1),
              Paragraph("一、调查概览", body), Paragraph("二、统计图表", body),
              Paragraph("三、统计结果", body), Paragraph("四、AI 分析报告", body),
              Paragraph("五、建议与风险提示", body), PageBreak()]

    # ---- 一、概览 ----
    story.append(Paragraph("一、调查概览", h1))
    kpi = [["累计问卷", "完成率", "平均时长", "今日新增"],
           [str(s["total"]), f"{s['completion_rate']*100:.1f}%", s["avg_duration_text"], f"+{s['today_new']}"]]
    t = Table(kpi, colWidths=[4 * cm] * 4)
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), FONT), ("FONTSIZE", (0, 0), (-1, -1), 12),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#c9d6f0")),
        ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)]))
    story += [t, Spacer(1, 0.6 * cm)]

    # ---- 二、统计图表 ----
    story.append(Paragraph("二、统计图表", h1))
    chart_order = ["pie_0", "pie_1", "rating_bar", "multi_rank", "trend"]
    for name in chart_order:
        if name in charts:
            img = _fit_image(Image, str(charts[name]), max_w=15 * cm)
            story += [img, Spacer(1, 0.4 * cm)]
    story.append(PageBreak())

    # ---- 三、统计结果 ----
    story.append(Paragraph("三、统计结果", h1))
    for st in analyzer.all_question_stats():
        story.append(Paragraph(f"{st['id']}. {st['title']}", h2))
        if st.get("type") in OPTION_TYPES and st.get("distribution"):
            data = [["选项", "计数", "占比"]] + [[d["label"], str(d["count"]), f"{d['ratio']*100:.1f}%"]
                                              for d in st["distribution"]]
            story.append(_simple_table(Table, TableStyle, colors, FONT, data))
        elif st.get("type") in NUMERIC_TYPES and st.get("n"):
            story.append(Paragraph(f"样本 {st['n']}；均值 {st['mean']}，中位数 {st['median']}，"
                                  f"标准差 {st['std']}，范围 {st['min']}~{st['max']}。", body))
        elif st.get("type") == "text":
            story.append(Paragraph(f"共填写 {st.get('filled',0)} 条。示例：" +
                                  "；".join(st.get("samples", [])[:3]), body))
        story.append(Spacer(1, 0.3 * cm))
    story.append(PageBreak())

    # ---- 四、AI 报告 ----
    story.append(Paragraph("四、AI 分析报告", h1))
    for label, key in [("调查结论", "conclusions"), ("用户画像", "persona"), ("趋势预测", "trends")]:
        story.append(Paragraph(label, h2))
        for item in ai_report.get(key, []):
            story.append(Paragraph("• " + item, body))
    story.append(PageBreak())

    # ---- 五、建议与风险 ----
    story.append(Paragraph("五、建议与风险提示", h1))
    for label, key in [("改进建议", "suggestions"), ("风险提示", "risks")]:
        story.append(Paragraph(label, h2))
        for item in ai_report.get(key, []):
            story.append(Paragraph("• " + item, body))
    story.append(Paragraph("智能洞察（精选）", h2))
    for i, ins in enumerate(ai_report.get("insights", [])[:10], 1):
        story.append(Paragraph(f"{i}. {ins}", body))

    out = _exports_dir(config, base_dir) / f"survey_report_{_ts()}.pdf"
    SimpleDocTemplate(str(out), pagesize=A4, title="适老化智能问卷调查分析报告").build(story)
    log.info("已导出 PDF：%s", out)
    return out


def _fit_image(Image, path: str, max_w: float):
    from PIL import Image as PILImage  # reportlab 依赖 pillow
    try:
        iw, ih = PILImage.open(path).size
        ratio = ih / iw
        return Image(path, width=max_w, height=max_w * ratio)
    except Exception:  # noqa: BLE001
        return Image(path, width=max_w, height=max_w * 0.6)


def _simple_table(Table, TableStyle, colors, font, data):
    t = Table(data, hAlign="LEFT")
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font), ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbe4f5")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9d6f0")),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
    return t


# ======================================================================
# Word 导出
# ======================================================================
def export_word(config: Dict[str, Any], q: Questionnaire, analyzer: Analyzer,
                base_dir: Path = BASE_DIR, ai_report: Optional[Dict[str, Any]] = None) -> Path:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    analyzer.refresh()
    ai_report = ai_report or build_ai_report(config, q, analyzer)
    charts = render_charts(q, analyzer, _exports_dir(config, base_dir) / f"_charts_{_ts()}")
    s = ai_report["summary"]
    doc = Document()

    # 封面
    title = doc.add_heading(config.get("export", {}).get("report_title", "适老化智能问卷调查分析报告"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(q.title()); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(f"样本规模：{s['total']} 份　|　完成率：{s['completion_rate']*100:.1f}%　|　"
                             f"生成时间：{ai_report['generated_at']}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 一、概览
    doc.add_heading("一、调查概览", level=1)
    table = doc.add_table(rows=2, cols=4); table.style = "Light Grid Accent 1"
    hdr = ["累计问卷", "完成率", "平均时长", "今日新增"]
    vals = [str(s["total"]), f"{s['completion_rate']*100:.1f}%", s["avg_duration_text"], f"+{s['today_new']}"]
    for i, (htxt, vtxt) in enumerate(zip(hdr, vals)):
        table.cell(0, i).text = htxt
        table.cell(1, i).text = vtxt

    # 二、统计图表
    doc.add_heading("二、统计图表", level=1)
    for name in ["pie_0", "pie_1", "rating_bar", "multi_rank", "trend"]:
        if name in charts:
            doc.add_picture(str(charts[name]), width=Inches(5.2))

    # 三、统计结果
    doc.add_heading("三、统计结果", level=1)
    for st in analyzer.all_question_stats():
        doc.add_heading(f"{st['id']}. {st['title']}", level=2)
        if st.get("type") in OPTION_TYPES and st.get("distribution"):
            t = doc.add_table(rows=1, cols=3); t.style = "Light List Accent 1"
            for j, htxt in enumerate(["选项", "计数", "占比"]):
                t.cell(0, j).text = htxt
            for d in st["distribution"]:
                row = t.add_row().cells
                row[0].text, row[1].text, row[2].text = d["label"], str(d["count"]), f"{d['ratio']*100:.1f}%"
        elif st.get("type") in NUMERIC_TYPES and st.get("n"):
            doc.add_paragraph(f"样本 {st['n']}；均值 {st['mean']}，中位数 {st['median']}，"
                              f"标准差 {st['std']}，范围 {st['min']}~{st['max']}。")
        elif st.get("type") == "text":
            doc.add_paragraph(f"共填写 {st.get('filled',0)} 条。示例：" + "；".join(st.get("samples", [])[:3]))

    # 四、AI 报告 + 五、建议风险
    for heading, keys in [("四、AI 分析报告", [("调查结论", "conclusions"), ("用户画像", "persona"), ("趋势预测", "trends")]),
                          ("五、建议与风险提示", [("改进建议", "suggestions"), ("风险提示", "risks")])]:
        doc.add_heading(heading, level=1)
        for label, key in keys:
            doc.add_heading(label, level=2)
            for item in ai_report.get(key, []):
                doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("智能洞察（精选）", level=2)
    for ins in ai_report.get("insights", [])[:10]:
        doc.add_paragraph(ins, style="List Number")

    out = _exports_dir(config, base_dir) / f"survey_report_{_ts()}.docx"
    doc.save(out)
    log.info("已导出 Word：%s", out)
    return out


# ======================================================================
# 示例数据生成
# ======================================================================
def generate_sample_data(config: Dict[str, Any], q: Optional[Questionnaire] = None,
                         n: int = 200, reset: bool = False, base_dir: Path = BASE_DIR) -> int:
    """生成 n 份模拟答卷（默认 200）。reset=True 时先清空已有数据。"""
    q = q or Questionnaire.load(config=config, base_dir=base_dir)
    mgr = dm.DataManager(config=config, base_dir=base_dir,
                         question_ids=[x.get("id") for x in q.questions])
    if reset:
        mgr.clear_all()
    return mgr.simulate_responses(q, n=n)


# ======================================================================
# 报告操作面板
# ======================================================================
class ReportPanel:
    """AI 报告与示例数据图形面板。"""

    def __init__(self, config: Dict[str, Any], master: Optional[tk.Misc] = None,
                 base_dir: Path = BASE_DIR) -> None:
        self.config = config
        self.base_dir = base_dir
        self.q = Questionnaire.load(config=config, base_dir=base_dir)
        self.data = dm.DataManager(config=config, base_dir=base_dir,
                                   question_ids=[x.get("id") for x in self.q.questions])
        self.analyzer = Analyzer(self.q, self.data)

        if master is not None:
            self.win: tk.Misc = tk.Toplevel(master)
            self.owns_loop = False
        else:
            self.win = tk.Tk()
            self.owns_loop = True
        self.win.title("AI 分析与报告导出")
        self.win.geometry("860x680")
        self.win.configure(bg="#f4f6fb")
        self._build()
        self._refresh_count()

    def _build(self) -> None:
        tk.Label(self.win, text="🤖 AI 分析与报告导出", font=("Microsoft YaHei", 20, "bold"),
                 bg="#0d1b2a", fg="#e0e1dd").pack(fill="x", ipady=14)
        bar = tk.Frame(self.win, bg="#e6edfb")
        bar.pack(fill="x")
        self.count_lbl = tk.Label(bar, text="", font=("Microsoft YaHei", 12), bg="#e6edfb")
        self.count_lbl.pack(side="left", padx=14, pady=8)

        def _darken(hexc: str, f: float = 0.85) -> str:
            r, g, bl = (int(hexc[i:i + 2], 16) for i in (1, 3, 5))
            return "#%02x%02x%02x" % (int(r * f), int(g * f), int(bl * f))

        def b(parent, txt, cmd, color="#2563eb"):
            # 用 Label 实现按钮：macOS 上原生 tk.Button 会忽略背景色，导致白字落在
            # 系统浅色按钮面上看不清。Label 完全遵守 bg/fg，彩色底 + 白字清晰可读。
            btn = tk.Label(parent, text=txt, font=("Microsoft YaHei", 12, "bold"),
                           bg=color, fg="white", padx=16, pady=9, cursor="hand2")
            btn.bind("<Button-1>", lambda e: cmd())
            btn.bind("<Enter>", lambda e: btn.configure(bg=_darken(color)))
            btn.bind("<Leave>", lambda e: btn.configure(bg=color))
            return btn

        b(bar, "🎲 生成200份示例数据", self.on_sample, "#475569").pack(side="right", padx=6, pady=6)

        actions = tk.Frame(self.win, bg="#f4f6fb")
        actions.pack(fill="x", pady=10, padx=14)
        b(actions, "🧠 生成AI报告预览", self.on_preview, "#0ea5a4").pack(side="left", padx=6)
        b(actions, "📊 导出 Excel", self.on_excel).pack(side="left", padx=6)
        b(actions, "📄 导出 PDF", self.on_pdf, "#e11d48").pack(side="left", padx=6)
        b(actions, "📝 导出 Word", self.on_word, "#2563eb").pack(side="left", padx=6)

        self.out = scrolledtext.ScrolledText(self.win, font=("Microsoft YaHei", 11), wrap="word")
        self.out.pack(fill="both", expand=True, padx=14, pady=10)
        self.status = tk.Label(self.win, text="就绪", font=("Microsoft YaHei", 10),
                               bg="#0d1b2a", fg="#9fb3d1", anchor="w")
        self.status.pack(fill="x", ipady=4)

    def _refresh_count(self) -> None:
        self.count_lbl.configure(text=f"当前答卷：{self.data.count()} 份")

    def _log(self, text: str) -> None:
        self.out.insert("end", text + "\n")
        self.out.see("end")
        self.win.update_idletasks()

    def _set_status(self, text: str) -> None:
        self.status.configure(text=text)
        self.win.update_idletasks()

    def on_sample(self) -> None:
        if not messagebox.askyesno("生成示例数据", "将生成 200 份模拟答卷并写入数据库，用于展示。是否继续？"):
            return
        self._set_status("正在生成示例数据…")
        n = generate_sample_data(self.config, self.q, n=200, base_dir=self.base_dir)
        self.analyzer.refresh()
        self._refresh_count()
        self._log(f"✅ 已生成 {n} 份模拟答卷。")
        self._set_status("示例数据生成完成")

    def on_preview(self) -> None:
        self._set_status("正在生成 AI 报告…")
        self.analyzer.refresh()
        rep = build_ai_report(self.config, self.q, self.analyzer)
        self.out.delete("1.0", "end")
        self._log(f"【{rep['title']}】  模式：{'联网AI' if rep['mode']=='online' else '离线规则'}\n")
        for label, key in [("调查结论", "conclusions"), ("用户画像", "persona"), ("趋势预测", "trends"),
                           ("改进建议", "suggestions"), ("风险提示", "risks")]:
            self._log(f"— {label} —")
            for item in rep.get(key, []):
                self._log("  • " + item)
            self._log("")
        self._log("— 智能洞察 —")
        for i, ins in enumerate(rep.get("insights", []), 1):
            self._log(f"  {i}. {ins}")
        self._set_status("AI 报告已生成")

    def _export(self, fn, kind: str) -> None:
        if self.data.count() == 0:
            messagebox.showinfo("提示", "暂无数据，请先生成示例数据或采集答卷。")
            return
        self._set_status(f"正在导出 {kind}…")
        try:
            path = fn(self.config, self.q, self.analyzer, base_dir=self.base_dir)
            self._log(f"✅ {kind} 导出成功：{path}")
            self._set_status(f"{kind} 导出完成")
            messagebox.showinfo("导出成功", f"{kind} 已导出到：\n{path}")
        except Exception as exc:  # noqa: BLE001
            log.exception("%s 导出失败", kind)
            self._log(f"❌ {kind} 导出失败：{exc}")
            self._set_status(f"{kind} 导出失败")
            messagebox.showerror("导出失败", str(exc))

    def on_excel(self) -> None:
        self._export(export_excel, "Excel")

    def on_pdf(self) -> None:
        self._export(export_pdf, "PDF")

    def on_word(self) -> None:
        self._export(export_word, "Word")

    def run(self) -> None:
        if self.owns_loop:
            self.win.mainloop()


# ======================================================================
# 入口
# ======================================================================
def launch(config: Dict[str, Any], master: Optional[tk.Misc] = None) -> Optional[ReportPanel]:
    app = ReportPanel(config, master=master)
    app.run()
    return app
